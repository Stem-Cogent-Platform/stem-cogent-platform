from __future__ import annotations

import csv
import hashlib
import io
import json
import re
from dataclasses import dataclass, replace
from datetime import UTC, datetime
from email.utils import parsedate_to_datetime
from html.parser import HTMLParser
from pathlib import PurePosixPath
from typing import Any
from urllib.parse import urljoin

from defusedxml import ElementTree  # type: ignore[import-untyped]
from docx import Document
from pypdf import PdfReader


_API_RECORD_LIMIT = 500


@dataclass(frozen=True)
class NormalizedDocument:
    signal_type: str
    title: str | None
    body_text: str
    source_url: str
    published_at: datetime | None
    original_language: str
    region_tags: tuple[str, ...]
    processing_flags: tuple[str, ...]

    @property
    def body_text_hash(self) -> str:
        return f"sha256:{hashlib.sha256(self.body_text.encode('utf-8')).hexdigest()}"


def normalize_payload(
    source_type: str,
    body: bytes,
    source_url: str,
    *,
    content_type: str = "application/octet-stream",
) -> tuple[NormalizedDocument, ...]:
    if not body:
        raise ValueError("Cannot normalize an empty payload")
    parsers = {
        "RSS": _rss_documents,
        "API": _api_documents,
        "HTML": _html_documents,
        "PDF": _pdf_documents,
        "USER_UPLOAD": _upload_documents,
        "LIVE_SEARCH": _discovery_documents,
    }
    try:
        documents = parsers[source_type](body, source_url, content_type)
    except KeyError as exc:
        raise ValueError(f"Unsupported normalization source type: {source_type}") from exc
    if not documents:
        raise ValueError("Payload did not contain a normalizable document")
    return tuple(documents)


def _document(
    signal_type: str,
    title: str | None,
    body: str,
    source_url: str,
    published_at: datetime | None = None,
    *,
    flags: tuple[str, ...] = (),
) -> NormalizedDocument:
    cleaned = _clean_text(body)
    if not cleaned:
        raise ValueError("Parsed document contains no text")
    return NormalizedDocument(
        signal_type=signal_type,
        title=_clean_text(title) if title else None,
        body_text=cleaned,
        source_url=source_url,
        published_at=published_at,
        original_language="en",
        region_tags=("NG",),
        processing_flags=flags,
    )


def _rss_documents(body: bytes, source_url: str, _: str) -> list[NormalizedDocument]:
    root = ElementTree.fromstring(body)
    items = root.findall(".//item") or root.findall(".//{*}entry")
    documents: list[NormalizedDocument] = []
    for item in items:
        title = _xml_text(item, "title")
        description = _xml_text(item, "description") or _xml_text(item, "summary")
        link = urljoin(source_url, _xml_link(item) or source_url)
        published = _parse_datetime(
            _xml_text(item, "pubDate")
            or _xml_text(item, "published")
            or _xml_text(item, "updated")
            or _xml_text(item, "documentDate")
        )
        documents.append(_document("FEED_ITEM", title, description or title or "", link, published))
    return documents


def _xml_text(element: Any, local_name: str) -> str | None:
    child = element.find(local_name)
    if child is None:
        child = element.find(f"{{*}}{local_name}")
    return child.text if child is not None and child.text else None


def _xml_link(element: Any) -> str | None:
    child = element.find("link")
    if child is None:
        child = element.find("{*}link")
    if child is None:
        return None
    return child.get("href") or child.text


def _api_documents(body: bytes, source_url: str, _: str) -> list[NormalizedDocument]:
    payload = json.loads(body)
    records = payload if isinstance(payload, list) else [payload]
    window_limited = len(records) > _API_RECORD_LIMIT
    documents: list[NormalizedDocument] = []
    for index, value in enumerate(records[:_API_RECORD_LIMIT]):
        record = value if isinstance(value, dict) else {"value": value}
        title = str(
            record.get("title")
            or record.get("name")
            or record.get("service")
            or f"API record {index + 1}"
        )
        published = _parse_datetime(
            str(
                record.get("published_at")
                or record.get("started_at")
                or record.get("timestamp")
                or record.get("documentDate")
                or ""
            )
        )
        record_url = urljoin(
            source_url,
            str(record.get("url") or record.get("link") or source_url),
        )
        text = "\n".join(f"{key}: {_scalar(value)}" for key, value in sorted(record.items()))
        documents.append(
            _document(
                "API_RECORD",
                title,
                text,
                record_url,
                published,
                flags=("LATEST_RECORD_WINDOW",) if window_limited else (),
            )
        )
    return documents


def _discovery_documents(body: bytes, source_url: str, _: str) -> list[NormalizedDocument]:
    payload = json.loads(body)
    articles = payload.get("articles", []) if isinstance(payload, dict) else []
    documents: list[NormalizedDocument] = []
    for article in articles:
        if not isinstance(article, dict) or not article.get("url"):
            continue
        title = str(article.get("title") or article["url"])
        metadata = {
            key: article[key]
            for key in ("domain", "language", "sourcecountry")
            if article.get(key)
        }
        detail = "\n".join(
            (title, *(f"{key}: {_scalar(value)}" for key, value in sorted(metadata.items())))
        )
        documents.append(
            _document(
                "DISCOVERED_ARTICLE",
                title,
                detail,
                str(article["url"]),
                _parse_datetime(str(article.get("seendate") or "")),
                flags=("DISCOVERY_LEAD", "REQUIRES_CORROBORATION"),
            )
        )
    return documents


class _VisibleHTMLParser(HTMLParser):
    def __init__(self) -> None:
        super().__init__(convert_charrefs=True)
        self.title: list[str] = []
        self.text: list[str] = []
        self.published_at: str | None = None
        self.canonical_url: str | None = None
        self.language = "en"
        self._ignored_depth = 0
        self._in_heading = False

    def handle_starttag(self, tag: str, attrs: list[tuple[str, str | None]]) -> None:
        attributes = dict(attrs)
        if tag in {"script", "style", "noscript", "svg"}:
            self._ignored_depth += 1
        if tag in {"title", "h1"}:
            self._in_heading = True
        if tag == "time" and attributes.get("datetime"):
            self.published_at = attributes["datetime"]
        if tag == "link" and attributes.get("rel") == "canonical":
            self.canonical_url = attributes.get("href")
        language = attributes.get("lang")
        if tag == "html" and language:
            self.language = language.split("-", maxsplit=1)[0]

    def handle_endtag(self, tag: str) -> None:
        if tag in {"script", "style", "noscript", "svg"} and self._ignored_depth:
            self._ignored_depth -= 1
        if tag in {"title", "h1"}:
            self._in_heading = False

    def handle_data(self, data: str) -> None:
        if self._ignored_depth:
            return
        self.text.append(data)
        if self._in_heading:
            self.title.append(data)


def _html_documents(body: bytes, source_url: str, _: str) -> list[NormalizedDocument]:
    parser = _VisibleHTMLParser()
    parser.feed(body.decode("utf-8", errors="replace"))
    url = urljoin(source_url, parser.canonical_url) if parser.canonical_url else source_url
    document = _document(
        "WEB_DOCUMENT",
        " ".join(parser.title),
        " ".join(parser.text),
        url,
        _parse_datetime(parser.published_at),
    )
    return [replace(document, original_language=parser.language)]


def _pdf_documents(body: bytes, source_url: str, _: str) -> list[NormalizedDocument]:
    # Public-sector PDFs frequently contain recoverable xref/stream defects.
    # Evidence integrity is established before parsing, so tolerant extraction is safe.
    reader = PdfReader(io.BytesIO(body), strict=False)
    text = "\n".join(page.extract_text() or "" for page in reader.pages)
    title = reader.metadata.title if reader.metadata else None
    return [_document("PDF_DOCUMENT", title, text, source_url)]


def _upload_documents(body: bytes, source_url: str, content_type: str) -> list[NormalizedDocument]:
    suffix = PurePosixPath(source_url).suffix.lower()
    if suffix == ".csv" or "csv" in content_type:
        rows = list(csv.DictReader(io.StringIO(body.decode("utf-8-sig"))))
        return [
            _document(
                "UPLOAD_RECORD",
                f"Upload row {index}",
                "\n".join(f"{key}: {value}" for key, value in row.items()),
                f"{source_url}#row={index}",
                flags=("PROPRIETARY_UPLOAD",),
            )
            for index, row in enumerate(rows, start=2)
        ]
    if suffix == ".docx" or "wordprocessingml" in content_type:
        document = Document(io.BytesIO(body))
        text = "\n".join(paragraph.text for paragraph in document.paragraphs)
        return [_document("UPLOAD_DOCUMENT", None, text, source_url, flags=("PROPRIETARY_UPLOAD",))]
    if suffix == ".pdf" or content_type == "application/pdf":
        return [replace(_pdf_documents(body, source_url, content_type)[0], processing_flags=("PROPRIETARY_UPLOAD",))]
    raise ValueError("Unsupported user-upload format")


def _parse_datetime(value: str | None) -> datetime | None:
    if not value:
        return None
    try:
        parsed = datetime.fromisoformat(value.replace("Z", "+00:00"))
    except ValueError:
        try:
            if re.fullmatch(r"\d{8}T\d{6}Z", value):
                parsed = datetime.strptime(value, "%Y%m%dT%H%M%SZ").replace(tzinfo=UTC)
            else:
                parsed = parsedate_to_datetime(value)
        except (TypeError, ValueError):
            try:
                parsed = datetime.strptime(value, "%d/%m/%Y").replace(tzinfo=UTC)
            except ValueError:
                return None
    if parsed.tzinfo is None:
        parsed = parsed.replace(tzinfo=UTC)
    return parsed.astimezone(UTC)


def _scalar(value: Any) -> str:
    if isinstance(value, (dict, list)):
        return json.dumps(value, sort_keys=True, separators=(",", ":"))
    return str(value)


def _clean_text(value: str | None) -> str:
    return re.sub(r"\s+", " ", value or "").strip()
